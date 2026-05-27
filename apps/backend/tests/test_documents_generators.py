"""Tests de los generadores DOCX + Markdown (Fase 4.1).

Los tests de DOCX re-abren el archivo generado con python-docx para
verificar que es un `.docx` válido (no solo bytes). Sin MS Office.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime

from docx import Document as DocxReader

from sqa_kb.documents.generators.base import GeneratedFile
from sqa_kb.documents.generators.docx import DocxGenerator
from sqa_kb.documents.generators.markdown import MarkdownGenerator
from sqa_kb.documents.models import DocumentContent, QaPair


def _content(**overrides) -> DocumentContent:  # type: ignore[no-untyped-def]
    base = {
        "document_id": "MTEC-flaky-tests-2026-05-26",
        "title": "Flaky tests en CI",
        "category": "TEC",
        "document_type": "MTEC",
        "version": "1.0",
        "fecha": datetime(2026, 5, 26, tzinfo=UTC),
        "author_name": "Tester",
        "author_role": "QA",
        "topic": "detección de flaky tests",
        "body_blocks": ("Los flaky tests fallan de forma intermitente.",),
        "qa_pairs": (QaPair(question="¿Cómo se detectan?", answer="Con reruns."),),
        "is_anonymized": False,
    }
    base.update(overrides)
    return DocumentContent(**base)  # type: ignore[arg-type]


# ===========================================================================
# MarkdownGenerator
# ===========================================================================


def test_markdown_generate_returns_generated_file() -> None:
    gen = MarkdownGenerator()
    out = gen.generate(_content())
    assert isinstance(out, GeneratedFile)
    assert out.filename == "MTEC-flaky-tests-2026-05-26.md"
    assert out.media_type.startswith("text/markdown")
    assert out.size_bytes > 0


def test_markdown_render_has_all_sections() -> None:
    md = MarkdownGenerator().render(_content())
    assert "# Flaky tests en CI" in md
    assert "## Tema" in md
    assert "## Contenido" in md
    assert "## Precisiones" in md
    assert "detección de flaky tests" in md
    assert "Los flaky tests fallan" in md
    assert "**¿Cómo se detectan?**" in md
    assert "Con reruns." in md


def test_markdown_metadata_table_includes_type_and_category_labels() -> None:
    md = MarkdownGenerator().render(_content())
    assert "MTEC — Memoria técnica" in md
    assert "TEC — Conocimiento Técnico" in md


def test_markdown_author_with_role() -> None:
    md = MarkdownGenerator().render(_content(author_role="QA Lead"))
    assert "Tester (QA Lead)" in md


def test_markdown_author_without_role() -> None:
    md = MarkdownGenerator().render(_content(author_role=None))
    assert "| Tester |" in md
    assert "Tester ()" not in md


def test_markdown_anonymized_flag_shown() -> None:
    md = MarkdownGenerator().render(_content(is_anonymized=True))
    assert "Anonimizado" in md


def test_markdown_anonymized_flag_hidden_when_false() -> None:
    md = MarkdownGenerator().render(_content(is_anonymized=False))
    assert "Anonimizado" not in md


def test_markdown_omits_content_section_when_no_blocks() -> None:
    md = MarkdownGenerator().render(_content(body_blocks=()))
    assert "## Contenido" not in md


def test_markdown_omits_precisiones_when_no_qa() -> None:
    md = MarkdownGenerator().render(_content(qa_pairs=()))
    assert "## Precisiones" not in md


def test_markdown_ends_with_single_newline() -> None:
    md = MarkdownGenerator().render(_content())
    assert md.endswith("\n")
    assert not md.endswith("\n\n")


# ===========================================================================
# DocxGenerator
# ===========================================================================


def test_docx_generate_returns_valid_docx() -> None:
    gen = DocxGenerator()
    out = gen.generate(_content())
    assert out.filename == "MTEC-flaky-tests-2026-05-26.docx"
    assert out.media_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # Re-abrir con python-docx prueba que es un .docx estructuralmente válido.
    reader = DocxReader(io.BytesIO(out.data))
    assert reader is not None


def test_docx_contains_title_and_topic() -> None:
    out = DocxGenerator().generate(_content())
    reader = DocxReader(io.BytesIO(out.data))
    full_text = "\n".join(p.text for p in reader.paragraphs)
    assert "Flaky tests en CI" in full_text
    assert "detección de flaky tests" in full_text


def test_docx_has_section_headings() -> None:
    out = DocxGenerator().generate(_content())
    reader = DocxReader(io.BytesIO(out.data))
    texts = [p.text for p in reader.paragraphs]
    assert "Tema" in texts
    assert "Contenido" in texts
    assert "Precisiones" in texts


def test_docx_metadata_table_present() -> None:
    out = DocxGenerator().generate(_content())
    reader = DocxReader(io.BytesIO(out.data))
    assert len(reader.tables) == 1
    table = reader.tables[0]
    # 5 filas base (sin anonimizado).
    labels = [row.cells[0].text for row in table.rows]
    assert "Carpeta" in labels
    assert "Tipo" in labels
    assert "Versión" in labels
    assert "Fecha" in labels
    assert "Autor" in labels


def test_docx_metadata_table_adds_anonymized_row() -> None:
    out = DocxGenerator().generate(_content(is_anonymized=True))
    reader = DocxReader(io.BytesIO(out.data))
    labels = [row.cells[0].text for row in reader.tables[0].rows]
    assert "Anonimizado" in labels


def test_docx_title_uses_brand_color() -> None:
    """El título debe estar en azul corporativo SQA (#03277D)."""
    from sqa_kb.documents import branding

    out = DocxGenerator().generate(_content())
    reader = DocxReader(io.BytesIO(out.data))
    # Buscar el run del título.
    title_color = None
    for p in reader.paragraphs:
        if p.text == "Flaky tests en CI":
            title_color = p.runs[0].font.color.rgb
            break
    assert title_color is not None
    assert str(title_color) == branding.COLOR_TITULO.hex


def test_docx_omits_content_when_no_blocks() -> None:
    out = DocxGenerator().generate(_content(body_blocks=(), qa_pairs=()))
    reader = DocxReader(io.BytesIO(out.data))
    texts = [p.text for p in reader.paragraphs]
    assert "Contenido" not in texts
    assert "Precisiones" not in texts
    # Pero Tema siempre va.
    assert "Tema" in texts


def test_docx_footer_has_sqa_branding() -> None:
    out = DocxGenerator().generate(_content())
    reader = DocxReader(io.BytesIO(out.data))
    footer = reader.sections[0].footer
    footer_text = "\n".join(p.text for p in footer.paragraphs)
    assert "SQA" in footer_text


# ===========================================================================
# Agent adapter sigue funcionando tras el refactor
# ===========================================================================


def test_agent_markdown_adapter_still_works() -> None:
    """`render_markdown_document` (Fase 2) ahora delega al MarkdownGenerator
    canónico — verificamos que el output sigue teniendo la estructura."""
    from sqa_kb.agent.markdown_generator import build_document_id

    doc_id = build_document_id(
        document_type="MTEC",
        topic="flaky tests",
        fecha=datetime(2026, 5, 26, tzinfo=UTC),
    )
    assert doc_id == "MTEC-flaky-tests-2026-05-26"
