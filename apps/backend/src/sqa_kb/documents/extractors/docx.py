"""DocxExtractor — lee texto + estructura de un `.docx` (Fase 4.3).

Usa python-docx. Detecta secciones por estilo de heading (los párrafos
con `style.name` que empieza con "Heading"). Si el documento no tiene
headings, devuelve una sola sección con todo el texto.

También extrae el texto de las tablas (concatenado fila por fila) — útil
para los documentos generados por `DocxGenerator` que ponen la metadata
en una tabla.
"""

from __future__ import annotations

import io

from docx import Document

from sqa_kb.documents.extractors.base import ExtractedDocument, ExtractedSection

EXTENSIONS = ("docx",)


class DocxExtractor:
    """Implementa `DocumentExtractor` para `.docx`."""

    @property
    def extensions(self) -> tuple[str, ...]:
        return EXTENSIONS

    def extract(self, data: bytes) -> ExtractedDocument:
        doc = Document(io.BytesIO(data))

        sections: list[ExtractedSection] = []
        current_title = ""
        current_body: list[str] = []
        all_text: list[str] = []

        def _flush() -> None:
            if current_title or current_body:
                sections.append(
                    ExtractedSection(
                        title=current_title,
                        content="\n".join(current_body).strip(),
                    )
                )

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            all_text.append(text)
            style = (para.style.name or "") if para.style else ""
            if style.startswith("Heading") or style == "Title":
                _flush()
                current_title = text
                current_body = []
            else:
                current_body.append(text)
        _flush()

        # Texto de tablas (la metadata del DocxGenerator vive en tabla).
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    all_text.append(" | ".join(cells))

        return ExtractedDocument(
            text="\n".join(all_text).strip(),
            sections=tuple(sections),
            page_count=0,
        )
