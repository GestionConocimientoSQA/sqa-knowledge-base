"""DocxGenerator — documentos Word con branding SQA (Fase 4.1).

Produce un `.docx` con la estructura común de los documentos del KB:
portada (título + tipo + carpeta) → tabla de metadata → Tema →
Contenido → Precisiones → footer corporativo.

Branding aplicado:
- Título en Exo 2, color azul corporativo, con barra naranja debajo
  (la firma visual de SQA).
- Headings de sección en azul medio.
- Tabla de metadata con header azul corp + texto blanco.
- Footer con `SQA · Gestión del Conocimiento` + número de página.

Genera todo programáticamente (sin plantilla `.docx` binaria) — decisión
de Fase 4: el código es la fuente de verdad del branding, sin assets
binarios en el repo.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from sqa_kb.documents import branding
from sqa_kb.documents.branding import RgbColor
from sqa_kb.documents.doc_types import category_label, doc_type_label
from sqa_kb.documents.generators.base import GeneratedFile
from sqa_kb.documents.models import DocumentContent

MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
EXTENSION = "docx"


def _rgb(color: RgbColor) -> RGBColor:
    """Convierte el `RgbColor` del branding al `RGBColor` de python-docx."""
    return RGBColor.from_string(color.hex)


def _set_cell_background(cell, color: RgbColor) -> None:  # type: ignore[no-untyped-def]
    """Setea el shading (fondo) de una celda de tabla.

    python-docx no expone shading de celda directamente — hay que tocar
    el XML (`w:shd`). Patrón estándar documentado en la comunidad.
    """
    shd = cell._tc.get_or_add_tcPr().makeelement(
        qn("w:shd"), {qn("w:fill"): color.hex}
    )
    cell._tc.get_or_add_tcPr().append(shd)


class DocxGenerator:
    """Genera `.docx` con branding SQA. Implementa `DocumentGenerator`."""

    @property
    def extension(self) -> str:
        return EXTENSION

    @property
    def media_type(self) -> str:
        return MEDIA_TYPE

    def generate(self, content: DocumentContent) -> GeneratedFile:
        doc = Document()
        self._apply_base_styles(doc)
        self._add_title_block(doc, content)
        self._add_metadata_table(doc, content)
        self._add_topic(doc, content)
        self._add_body(doc, content)
        self._add_qa(doc, content)
        self._add_footer(doc)

        buffer = io.BytesIO()
        doc.save(buffer)
        return GeneratedFile(
            filename=f"{content.document_id}.{EXTENSION}",
            media_type=MEDIA_TYPE,
            data=buffer.getvalue(),
        )

    # -- estilos base -------------------------------------------------------

    def _apply_base_styles(self, doc: Document) -> None:  # type: ignore[no-untyped-def]
        """Setea la fuente de cuerpo por defecto del documento."""
        normal = doc.styles["Normal"]
        normal.font.name = branding.FONT_BODY
        normal.font.size = Pt(branding.SIZE_BODY)
        normal.font.color.rgb = _rgb(branding.COLOR_CUERPO)

    # -- portada ------------------------------------------------------------

    def _add_title_block(self, doc: Document, content: DocumentContent) -> None:  # type: ignore[no-untyped-def]
        # Línea superior: tipo + carpeta (caption).
        tipo = doc_type_label(content.document_type)
        carpeta = category_label(content.category)
        caption = doc.add_paragraph()
        run = caption.add_run(f"{tipo}  ·  {carpeta}")
        run.font.name = branding.FONT_BODY
        run.font.size = Pt(branding.SIZE_CAPTION)
        run.font.color.rgb = _rgb(branding.COLOR_TEXTO_SECUNDARIO)
        run.font.all_caps = True

        # Título principal.
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(content.title)
        title_run.font.name = branding.FONT_DISPLAY
        title_run.font.size = Pt(branding.SIZE_H1)
        title_run.font.bold = True
        title_run.font.color.rgb = _rgb(branding.COLOR_TITULO)

        # Barra naranja (firma visual SQA) — párrafo con borde inferior
        # grueso de color acento.
        bar = doc.add_paragraph()
        self._set_bottom_border(bar, branding.COLOR_ACENTO, size_pt=3)

    def _set_bottom_border(self, paragraph, color: RgbColor, *, size_pt: int) -> None:  # type: ignore[no-untyped-def]
        """Agrega un borde inferior de color al párrafo (la barra naranja)."""
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
        bottom = p_bdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): str(size_pt * 8),  # eighths of a point
                qn("w:space"): "1",
                qn("w:color"): color.hex,
            },
        )
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    # -- metadata -----------------------------------------------------------

    def _add_metadata_table(self, doc: Document, content: DocumentContent) -> None:  # type: ignore[no-untyped-def]
        autor = content.author_name
        if content.author_role:
            autor = f"{content.author_name} ({content.author_role})"
        rows = [
            ("Carpeta", f"{content.category} — {category_label(content.category)}"),
            ("Tipo", f"{content.document_type} — {doc_type_label(content.document_type)}"),
            ("Versión", content.version),
            ("Fecha", content.fecha.strftime("%Y-%m-%d")),
            ("Autor", autor),
        ]
        if content.is_anonymized:
            rows.append(("Anonimizado", "Sí"))

        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        for i, (label, value) in enumerate(rows):
            label_cell = table.rows[i].cells[0]
            value_cell = table.rows[i].cells[1]
            # Columna de label: fondo azul corp + texto blanco bold.
            _set_cell_background(label_cell, branding.COLOR_FONDO_TABLA_HEADER)
            self._style_cell(label_cell, label, color=branding.BLANCO, bold=True)
            self._style_cell(value_cell, value, color=branding.COLOR_CUERPO, bold=False)
        doc.add_paragraph()

    def _style_cell(self, cell, text: str, *, color: RgbColor, bold: bool) -> None:  # type: ignore[no-untyped-def]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.name = branding.FONT_BODY
        run.font.size = Pt(branding.SIZE_BODY)
        run.font.bold = bold
        run.font.color.rgb = _rgb(color)

    # -- secciones ----------------------------------------------------------

    def _add_section_heading(self, doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = branding.FONT_DISPLAY
        run.font.size = Pt(branding.SIZE_H2)
        run.font.bold = True
        run.font.color.rgb = _rgb(branding.COLOR_SUBTITULO)

    def _add_topic(self, doc: Document, content: DocumentContent) -> None:  # type: ignore[no-untyped-def]
        self._add_section_heading(doc, "Tema")
        doc.add_paragraph(content.topic)

    def _add_body(self, doc: Document, content: DocumentContent) -> None:  # type: ignore[no-untyped-def]
        if not content.body_blocks:
            return
        self._add_section_heading(doc, "Contenido")
        for block in content.body_blocks:
            doc.add_paragraph(block)

    def _add_qa(self, doc: Document, content: DocumentContent) -> None:  # type: ignore[no-untyped-def]
        if not content.qa_pairs:
            return
        self._add_section_heading(doc, "Precisiones")
        for qa in content.qa_pairs:
            q = doc.add_paragraph()
            q_run = q.add_run(qa.question)
            q_run.font.bold = True
            q_run.font.color.rgb = _rgb(branding.COLOR_CUERPO)
            doc.add_paragraph(qa.answer)

    # -- footer -------------------------------------------------------------

    def _add_footer(self, doc: Document) -> None:  # type: ignore[no-untyped-def]
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(branding.DOC_FOOTER_LEFT)
        run.font.name = branding.FONT_BODY
        run.font.size = Pt(branding.SIZE_FOOTER)
        run.font.color.rgb = _rgb(branding.COLOR_TEXTO_SECUNDARIO)
